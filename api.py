from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import platform
import subprocess
import base64

app = Flask(__name__)

# Path to store the serial number counter
SERIAL_FILE = "serial_data.txt"

def get_serial_number():
    # Read base number and counter from file
    with open(SERIAL_FILE, "r") as f:
        base_number, counter = map(int, f.read().strip().split(","))

    # Calculate current serial number
    serial_number = base_number + counter

    # Increment the counter and update the file
    with open(SERIAL_FILE, "w") as f:
        f.write(f"{base_number},{counter + 1}")

    return serial_number

def generate_reference_number(company_name="BKR"):
    """
    Generate the full reference number in the format: BKRMM-YYYY-CR<serial>.
    """
    current_month = datetime.now().strftime("%m")
    current_year = datetime.now().strftime("%Y")
    serial_number = get_serial_number()
    return f"{company_name}{current_month}-{current_year}-CR{serial_number}"

# Code to replace placeholder for VAT
def replace_placeholders_vat(doc, placeholders):
    """Replace placeholders in a Word document, maintaining original formatting."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a paragraph, preserving formatting."""
        for run in paragraph.runs:
            if key in run.text:
                # Replace placeholder text
                run.text = run.text.replace(key, value)
                # Retain original font style and size
                run.font.name = paragraph.style.font.name
                run.font.size = paragraph.style.font.size

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc

def replace_placeholders(doc, placeholders):
    """Replace placeholders in a Word document, including paragraphs and tables."""
    
    def replace_in_paragraph(paragraph, key, value):
        """Replace placeholders in a single paragraph, handling split runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            full_text = full_text.replace(key, value)
            for run in paragraph.runs:
                run.text = ""  # Clear all runs
            paragraph.runs[0].text = full_text  # Add the replaced text back

    def replace_in_cell(cell, placeholders):
        """Replace placeholders inside a table cell."""
        for para in cell.paragraphs:
            for key, value in placeholders.items():
                replace_in_paragraph(para, key, value)

    # Replace placeholders in all paragraphs
    for para in doc.paragraphs:
        for key, value in placeholders.items():
            replace_in_paragraph(para, key, value)

    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell, placeholders)

    return doc
  

def convert_to_pdf(doc_path, pdf_path):
    doc_path = os.path.abspath(doc_path)
    pdf_path = os.path.abspath(pdf_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"Word document not found at {doc_path}")

    # Use a temporary directory for the intermediate PDF file
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_pdf_path = os.path.join(temp_dir, "temp_output.pdf")

        # Step 1: Convert Word to PDF
        if platform.system() == "Windows":
            try:
                import comtypes.client
                import pythoncom
                pythoncom.CoInitialize()
                word = comtypes.client.CreateObject("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(doc_path)
                doc.SaveAs(temp_pdf_path, FileFormat=17)  # FileFormat=17 is for PDF
                doc.Close()
                word.Quit()
            except Exception as e:
                raise Exception(f"Error using COM on Windows: {e}")
        else:
            try:
                subprocess.run(
                    ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', temp_dir, doc_path],
                    check=True
                )
            except subprocess.CalledProcessError as e:
                raise Exception(f"Error using LibreOffice: {e}")

        # Step 2: Flatten the PDF (convert to image-based PDF)
        flatten_pdf(temp_pdf_path, pdf_path)

def flatten_pdf(input_pdf_path, output_pdf_path):
    """
    Converts each page of a PDF into an image and re-embeds it to create a flattened, non-editable PDF.
    """
    if not os.path.exists(input_pdf_path):
        raise FileNotFoundError(f"Input PDF file not found: {input_pdf_path}")

    doc = fitz.open(input_pdf_path)  # Open the original PDF
    writer = PdfWriter()

    with tempfile.TemporaryDirectory() as temp_dir:
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=300)  # Render page to an image with 300 DPI
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Save the image as a temporary PDF
            temp_page_path = os.path.join(temp_dir, f"temp_page_{page_num}.pdf")
            image.save(temp_page_path, "PDF")

            # Read the temporary PDF and add it to the writer
            reader = PdfReader(temp_page_path)
            writer.add_page(reader.pages[0])

    # Save the flattened PDF
    with open(output_pdf_path, "wb") as f:
        writer.write(f)

    print(f"Flattened PDF saved at: {output_pdf_path}")

def generate_unique_reference():
    """
    Generate a unique reference number based on the current date and time in the format:
    DDMMYYYYHHMMSS
    """
    now = datetime.now()
    return now.strftime("%d%m%Y%H%M%S")

@app.route('/generate-document', methods=['POST'])
def generate_document():
    """API endpoint to generate a document."""
    data = request.json
    template_type = data.get("template_type")
    placeholders = data.get("placeholders", {})

    try:
        reference_number = generate_reference_number()
        placeholders["<<Reference Number>>"] = reference_number

        template_paths = {
            "VAT": "SAMPLE VAT registration and VAT filling -SME package.docx",
            "Service Agreement": "SAMPLE Service Agreement -Company formation -Bahrain - Filled.docx",
            "Invoice": "SAMPLE -Invoice BKR2024CF158 - first payment.docx"
        }

        template_path = template_paths.get(template_type)
        if not template_path or not os.path.exists(template_path):
            return jsonify({"status": "error", "message": "Invalid template type or template not found!"}), 400

        doc = Document(template_path)
        doc = replace_placeholders(doc, placeholders)

        word_output = f"{template_type} {reference_number}.docx"
        pdf_output = word_output.replace(".docx", ".pdf")

        doc.save(word_output)
        convert_to_pdf(word_output, pdf_output)

        return jsonify({
            "status": "success",
            "reference_number": reference_number,
            "word_document": word_output,
            "pdf_document": pdf_output
        })

    except Exception as e:
        return jsonify({"status": "error", "message": str(e)}), 500

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """API endpoint to download a generated document."""
    if os.path.exists(filename):
        return send_file(filename, as_attachment=True)
    else:
        return jsonify({"status": "error", "message": "File not found!"}), 404

@app.route("/", methods=["GET"])
def home():
    return jsonify({
        "message": "Document Generation API is running!",
        "endpoints": {
            "POST /generate-document": "Generate a document with placeholders",
            "GET /download/<filename>": "Download a generated document"
        }
    })

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
